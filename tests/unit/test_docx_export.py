"""Tests for the DOCX export of the Biblioteca de Referentes — isolated in-memory DB."""

from datetime import datetime, timezone

import pytest
from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from eci.database.models import Ad, Advertiser, Base
from eci.reports import docx_export, library as library_module


@pytest.fixture()
def session(tmp_path, monkeypatch):
    monkeypatch.setattr(library_module, "LIBRARY_DIR", tmp_path)
    monkeypatch.setattr(docx_export, "LIBRARY_DIR", tmp_path)
    engine = create_engine("sqlite:///:memory:", connect_args={"check_same_thread": False})
    Base.metadata.create_all(engine)
    SessionLocal = sessionmaker(bind=engine, expire_on_commit=False, future=True)
    s = SessionLocal()
    yield s
    s.close()


def _advertiser(session, page_id, page_name, *, ecommerce_score=85.0, subniche=None, country="CO", niche="TEXTIL"):
    adv = Advertiser(
        page_id=page_id, page_name=page_name, niche=niche, subniche=subniche, country=country,
        active_ad_count=5, store_url=f"https://{page_id}.com/product", ecommerce_score=ecommerce_score,
        shopify_detected=True, scale_signal_score=70.0, confidence_score=60.0,
        oldest_active_ad_age_days=94, last_verified_at=datetime.now(timezone.utc),
        dominant_hook="curiosidad", dominant_angle="ahorro", dominant_offer="envio_gratis",
        video_count=2, image_count=1, carousel_count=0,
    )
    session.add(adv)
    session.flush()
    return adv


def _ad(session, page_id, ad_id, age_days, *, hook_type="offer", angle="ahorro", offer="percentage_discount", format="image", claims_risk_flags=None, primary_text="Oferta especial de referencia"):
    ad = Ad(
        ad_id=ad_id, source_name="meta_web_scraper", page_id=page_id, page_name=page_id, active=True,
        age_days=age_days, format=format, hook_type=hook_type, creative_angle=angle, offer_type=offer,
        primary_text=primary_text, claims_risk_flags=claims_risk_flags or [],
        ad_library_url=f"https://www.facebook.com/ads/library/?id={ad_id}",
    )
    session.add(ad)
    session.flush()
    return ad


def test_build_docx_creates_a_valid_document_with_brand_content(session):
    _advertiser(session, "p1", "Tienda Word", subniche="vestidos")
    _ad(session, "p1", "a1", age_days=45)
    _ad(session, "p1", "a2", age_days=10)
    session.commit()

    path = docx_export.build_docx("TEXTIL", ["CO"], session)
    assert path.exists()
    assert path.suffix == ".docx"

    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            full_text += "\n" + " ".join(c.text for c in row.cells)

    assert "Tienda Word" in full_text
    assert "Todos los anuncios activos (2)" in full_text
    assert "curiosidad" in full_text


def test_build_docx_includes_claims_warning_for_sensitive_niche(session):
    _advertiser(session, "p1", "Salud Store", niche="SALUD")
    _ad(session, "p1", "a1", age_days=45, claims_risk_flags=["cura_enfermedad"])
    session.commit()

    path = docx_export.build_docx("SALUD", ["CO"], session)
    doc = Document(str(path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "claims de salud" in full_text.lower()


def test_hyperlinks_use_valid_relationship_and_schema_element_order(session):
    """Regression test: links didn't open in Word — root cause was `w:rPr` children in the
    wrong order (color before b/rStyle), which violates the OOXML schema (CT_RPr requires
    rStyle, then b, then color, then u) and makes Word flag the file for repair, silently
    stripping the malformed hyperlink runs. Also confirms every w:hyperlink's r:id actually
    resolves to a real External hyperlink relationship (not a dangling reference)."""
    import zipfile
    from xml.etree import ElementTree as ET

    _advertiser(session, "p1", "Tienda Word", subniche="vestidos")
    _ad(session, "p1", "a1", age_days=45)
    session.commit()

    path = docx_export.build_docx("TEXTIL", ["CO"], session)

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

    with zipfile.ZipFile(str(path)) as z:
        doc_xml = z.read("word/document.xml")
        rels_xml = z.read("word/_rels/document.xml.rels")

    rel_root = ET.fromstring(rels_xml)
    hyperlink_rel_ids = {
        el.get("Id")
        for el in rel_root
        if el.get("Type") == REL_TYPE and el.get("TargetMode") == "External" and el.get("Target")
    }
    assert hyperlink_rel_ids, "no hyperlink relationships were written at all"

    doc_root = ET.fromstring(doc_xml)
    hyperlinks = doc_root.iter(f"{W}hyperlink")
    checked = 0
    for hl in hyperlinks:
        r_id = hl.get(f"{R}id")
        assert r_id in hyperlink_rel_ids, f"w:hyperlink references a dangling r:id {r_id}"
        rpr = hl.find(f"{W}r/{W}rPr")
        assert rpr is not None, "hyperlink run has no rPr"
        tag_order = [child.tag.replace(W, "") for child in rpr]
        # CT_RPr schema order: rStyle, then (b), then color, then u — anything out of this
        # relative order is what made Word flag the file for repair.
        expected_subsequence = [t for t in ["rStyle", "b", "color", "u"] if t in tag_order]
        actual_relevant = [t for t in tag_order if t in ("rStyle", "b", "color", "u")]
        assert actual_relevant == expected_subsequence, f"rPr children out of schema order: {tag_order}"
        assert "rStyle" in tag_order, "hyperlink run doesn't reference the Hyperlink character style"
        checked += 1
    assert checked > 0, "no w:hyperlink elements were found to check"


def test_build_docx_handles_empty_niche_without_crashing(session):
    path = docx_export.build_docx("TEXTIL", ["CO"], session)
    assert path.exists()
    doc = Document(str(path))
    assert len(doc.paragraphs) > 0
