"""DOCX export of the Biblioteca de Referentes — one professional, visual Word document per
niche (operator, 2026-08-14: "necesito que se pueda descargar en un archivo Word muy bien
estructurado presentable y profesional... cada nicho por separado... que tenga los links,
las explicaciones, todo").

Renders from the exact same `_build_library_context()` the HTML/JSON outputs use, so all
three formats can never disagree about which brands or numbers are in a given build.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from eci.reports.library import LIBRARY_DIR, _build_library_context

_ACCENT = RGBColor(0x0E, 0x7A, 0x8C)
_MUTED = RGBColor(0x60, 0x68, 0x78)
_DARK = RGBColor(0x1A, 0x1D, 0x29)
_GOLD = RGBColor(0xA8, 0x6A, 0x14)
_DANGER = RGBColor(0xA8, 0x23, 0x18)
_TIER_COLORS = {"alta": RGBColor(0x1B, 0x8A, 0x5A), "media": RGBColor(0x1D, 0x5F, 0xA8), "emergente": RGBColor(0x60, 0x68, 0x78)}
_HEADER_SHADING = "0E7A8C"
_ALT_ROW_SHADING = "F2F6F7"


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _ensure_hyperlink_style(doc: Document) -> None:
    """python-docx's blank template has NO "Hyperlink" character style defined — confirmed
    by inspecting a fresh Document()'s styles collection. A `<w:hyperlink>` run that only
    carries direct color/underline formatting (no named style) is valid OOXML and Word
    itself tolerates it, but some converters/viewers (Google Docs import, some mobile/web
    Word renderers) key off the actual "Hyperlink" style to decide a run is a real,
    clickable link — without it, the link can render as plain unclickable text. Creating
    the style explicitly (once per document) is the standard, most-compatible fix."""
    from docx.enum.style import WD_STYLE_TYPE

    if "Hyperlink" in [s.name for s in doc.styles]:
        return
    style = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    style.font.color.rgb = _ACCENT
    style.font.underline = True


def _add_hyperlink(paragraph, url: str | None, text: str, *, color: RGBColor = _ACCENT, bold: bool = False):
    """python-docx has no built-in hyperlink helper — build the run's XML directly. Falls
    back to plain (non-linked) text if the URL is missing, so a card never silently loses
    its label just because a link wasn't available.

    Sets BOTH the named "Hyperlink" character style (w:rStyle, must be the first child of
    w:rPr per the OOXML schema — see _ensure_hyperlink_style) and direct color/underline
    formatting, so the link renders and is clickable in the widest range of viewers, not
    just desktop Word."""
    if not url:
        run = paragraph.add_run(text)
        run.font.color.rgb = _MUTED
        return run
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    if bold:
        b = OxmlElement("w:b")
        rpr.append(b)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rpr.append(color_el)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


_FORMAT_LABELS = {"video": "Video", "image": "Imagen", "carousel": "Carrusel", "unknown": "Sin determinar"}


def _fmt_label(fmt: str | None) -> str:
    return _FORMAT_LABELS.get(fmt or "unknown", fmt or "Sin determinar")


def build_docx(
    niche: str,
    markets: list[str],
    session,
    *,
    min_ad_age_days: int = 30,
    min_ecommerce_score: float = 70.0,
    top_n: int = 10,
    must_include_domains: list[str] | None = None,
) -> Path:
    ctx = _build_library_context(
        niche,
        markets,
        session,
        min_ad_age_days=min_ad_age_days,
        min_ecommerce_score=min_ecommerce_score,
        top_n=top_n,
        must_include_domains=must_include_domains,
    )
    niche = ctx["niche"]

    doc = Document()
    _ensure_hyperlink_style(doc)
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = _DARK

    # ---------- Cover ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("📚 ECI Suite — Biblioteca de Referentes")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _GOLD

    h = doc.add_heading(f"{niche.title()}", level=0)
    for r in h.runs:
        r.font.color.rgb = _DARK

    sub = doc.add_paragraph()
    sub.add_run(f"{' · '.join(markets)} · generado {ctx['generated_at'][:10]}").font.color.rgb = _MUTED

    crit = doc.add_paragraph()
    crit.add_run(
        "Solo entran marcas independientes (nunca SHEIN, Temu, Mercado Libre ni similares) con tienda "
        "ecommerce real verificada y al menos un anuncio activo con "
        f"{ctx['min_ad_age_days']}+ días corriendo — la señal de que ya encontraron una estrategia que "
        "les funciona, no solo un test. Se agrupan por nivel de escala: 50+ anuncios activos = escala "
        "alta; menos que eso, referencias más pequeñas pero igualmente reales y verificadas."
    ).font.color.rgb = _MUTED

    if ctx["is_claims_sensitive_niche"]:
        warn = doc.add_paragraph()
        wr = warn.add_run(
            "⚠️ Nicho sensible a claims de salud: los anuncios marcados con 🚩 usan lenguaje de riesgo "
            "(cura/trata/elimina/resultados garantizados). Se muestran como referencia de qué hace la "
            "competencia, no como copy recomendado para copiar."
        )
        wr.font.color.rgb = _DANGER
        wr.italic = True

    # ---------- Summary stats ----------
    doc.add_heading("Resumen", level=1)
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.style = "Light Grid Accent 1"
    hdr = stats_table.rows[0].cells
    for i, label in enumerate(["Marcas referentes", "Escala alta (50+)", "Subnichos", "Mercados"]):
        hdr[i].text = label
        _shade_cell(hdr[i], _HEADER_SHADING)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.bold = True
    row = stats_table.add_row().cells
    row[0].text = str(ctx["total"])
    row[1].text = str(ctx["tier_counts"].get("alta", 0))
    row[2].text = str(len(ctx["by_subniche"]))
    row[3].text = str(len(ctx["markets"]))

    # ---------- Format stats ----------
    doc.add_heading("¿Qué formato usan más?", level=2)
    fmt_p = doc.add_paragraph()
    fp = ctx["format_percentages"]
    fmt_p.add_run(
        f"🎬 Video · {fp.get('video', 0)}%   🖼️ Imagen · {fp.get('image', 0)}%   "
        f"🎠 Carrusel · {fp.get('carousel', 0)}%   ❔ Sin determinar · {fp.get('unknown', 0)}%"
    )
    fmt_p2 = doc.add_paragraph()
    fmt_p2.add_run(
        f"📌 {ctx['single_format_brands']} de {ctx['total']} marcas se enfocan en un solo formato · "
        f"🔀 {ctx['mixed_format_brands']} de {ctx['total']} combinan varios formatos"
    ).font.color.rgb = _MUTED

    if ctx["top_hooks"] or ctx["top_angles"] or ctx["top_offers"]:
        doc.add_heading("Patrones comunes (anuncios 30+ días)", level=2)
        for label, values in (("Hooks", ctx["top_hooks"]), ("Ángulos", ctx["top_angles"]), ("Ofertas", ctx["top_offers"])):
            if not values:
                continue
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(" · ".join(f"{v} ({c})" for v, c in values)).font.color.rgb = _MUTED

    doc.add_page_break()

    # ---------- Ranking ----------
    doc.add_heading(f"Ranking del nicho — {ctx['total']} marcas", level=1)

    for tier in ctx["tier_order"]:
        subniche_groups = ctx["by_tier_and_subniche"].get(tier, {})
        if not subniche_groups:
            continue
        tier_heading = doc.add_heading(ctx["tier_labels"][tier], level=2)
        for r in tier_heading.runs:
            r.font.color.rgb = _TIER_COLORS.get(tier, _DARK)

        for subniche, group in subniche_groups.items():
            doc.add_heading(f"{subniche} ({len(group)})", level=3)

            for b in group:
                # --- Brand heading ---
                bh = doc.add_paragraph()
                bh.paragraph_format.space_before = Pt(10)
                run = bh.add_run(f"#{b['rank']}  {b['page_name']}  {b['market_flag']}")
                run.bold = True
                run.font.size = Pt(12.5)
                run.font.color.rgb = _DARK
                tier_run = bh.add_run(f"   [{ctx['tier_labels'][b['scale_tier']].split(' (')[0]}]")
                tier_run.font.color.rgb = _TIER_COLORS.get(b["scale_tier"], _MUTED)
                tier_run.italic = True

                # --- Quick stats table ---
                info_table = doc.add_table(rows=1, cols=5)
                info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                info_table.style = "Light List Accent 1"
                cells = info_table.rows[0].cells
                labels = ["Anuncios activos", "Más antiguo", "Scale Score", "Confianza", "Shopify"]
                values = [
                    str(b["active_ad_count"]),
                    f"{b['oldest_active_ad_age_days']}d",
                    f"{b['scale_signal_score']}/100",
                    f"{b['confidence_score']}/100",
                    "Sí" if b["shopify_detected"] else "No verificado",
                ]
                for i in range(5):
                    c0 = cells[i]
                    p0 = c0.paragraphs[0]
                    p0.add_run(labels[i] + ": ").bold = True
                    p0.add_run(values[i])

                # --- Dominant pattern ---
                if b.get("dominant_hook") or b.get("dominant_angle") or (b.get("dominant_offer") and b["dominant_offer"] != "no_offer"):
                    dp = doc.add_paragraph()
                    dp.add_run("Oferta y ángulo dominante: ").bold = True
                    parts = []
                    if b.get("dominant_hook"):
                        parts.append(f"Hook: {b['dominant_hook']}")
                    if b.get("dominant_angle"):
                        parts.append(f"Ángulo: {b['dominant_angle']}")
                    if b.get("dominant_offer") and b["dominant_offer"] != "no_offer":
                        parts.append(f"Oferta: {b['dominant_offer']}")
                    dp.add_run(" · ".join(parts)).font.color.rgb = _MUTED

                # --- Variedad creativa ---
                vp = doc.add_paragraph()
                vp.add_run("Variedad creativa: ").bold = True
                vp.add_run(
                    f"🎬 {b.get('video_count', 0)} video · 🖼️ {b.get('image_count', 0)} imagen · "
                    f"🎠 {b.get('carousel_count', 0)} carrusel · ❔ {b.get('unknown_format_count', 0)} sin determinar"
                ).font.color.rgb = _MUTED

                # --- Links ---
                # Every link also gets its raw URL printed as plain, selectable/copy-able
                # text right below — belt-and-suspenders against any viewer that doesn't
                # render OOXML hyperlinks as clickable (confirmed independently, via the
                # mammoth library, that the hyperlinks themselves are valid — this is purely
                # so the link is still USABLE by copy-paste no matter what opens the file).
                lp = doc.add_paragraph()
                lp.add_run("🛍 ")
                _add_hyperlink(lp, b.get("store_url"), "Ver tienda", bold=True)
                lp.add_run("     📣 ")
                _add_hyperlink(
                    lp,
                    b.get("ads_library_page_url"),
                    "Ver anuncios en Meta Ad Library" if b.get("ads_library_is_page_level") else "Ver anuncios de la marca en Meta Ad Library",
                    bold=True,
                )
                if b.get("store_url"):
                    url_p = doc.add_paragraph()
                    url_p.paragraph_format.space_after = Pt(2)
                    r0 = url_p.add_run(f"Tienda: {b['store_url']}")
                    r0.font.size = Pt(8)
                    r0.font.color.rgb = _MUTED
                if b.get("ads_library_page_url"):
                    url_p2 = doc.add_paragraph()
                    r1 = url_p2.add_run(f"Meta Ad Library: {b['ads_library_page_url']}")
                    r1.font.size = Pt(8)
                    r1.font.color.rgb = _MUTED

                # --- Reference ads (30+ day, up to 2, with quote + structure note) ---
                if b.get("long_running_ads"):
                    ref_heading = doc.add_paragraph()
                    ref_heading.add_run("Anuncios de referencia (30+ días activos)").bold = True
                    for ad in b["long_running_ads"][:2]:
                        ap = doc.add_paragraph(style="List Bullet")
                        tags = [f"{ad['age_days']}d", _fmt_label(ad.get("format"))]
                        if ad.get("hook_type"):
                            tags.append(f"hook: {ad['hook_type']}")
                        if ad.get("creative_angle"):
                            tags.append(f"ángulo: {ad['creative_angle']}")
                        if ad.get("offer_type") and ad["offer_type"] != "no_offer":
                            tags.append(f"oferta: {ad['offer_type']}")
                        tag_run = ap.add_run(" · ".join(tags))
                        tag_run.font.size = Pt(9)
                        tag_run.font.color.rgb = _MUTED
                        if ad.get("claims_flags"):
                            flag_run = ap.add_run("  🚩 claim de riesgo")
                            flag_run.font.color.rgb = _DANGER
                            flag_run.bold = True
                        if ad.get("primary_text"):
                            qp = doc.add_paragraph()
                            qp.paragraph_format.left_indent = Cm(0.6)
                            qr = qp.add_run(f'"{ad["primary_text"]}"')
                            qr.italic = True
                        notes = b.get("creative_notes") or {}
                        if ad.get("format") and notes.get(ad["format"]):
                            np_ = doc.add_paragraph()
                            np_.paragraph_format.left_indent = Cm(0.6)
                            np_.add_run("🎨 Estructura: ").bold = True
                            np_.add_run(notes[ad["format"]]).font.color.rgb = _MUTED
                        link_p = doc.add_paragraph()
                        link_p.paragraph_format.left_indent = Cm(0.6)
                        _add_hyperlink(link_p, ad.get("ad_library_url"), "▶ Ver este anuncio en Meta Ad Library")
                        if ad.get("ad_library_url"):
                            raw_url_p = doc.add_paragraph()
                            raw_url_p.paragraph_format.left_indent = Cm(0.6)
                            raw_url_p.paragraph_format.space_after = Pt(2)
                            rr = raw_url_p.add_run(ad["ad_library_url"])
                            rr.font.size = Pt(7.5)
                            rr.font.color.rgb = _MUTED
                        credit_p = doc.add_paragraph()
                        credit_p.paragraph_format.left_indent = Cm(0.6)
                        cr = credit_p.add_run("Anuncio cortesía de Meta Ad Library — se abre en facebook.com, no se aloja en este documento.")
                        cr.font.size = Pt(8)
                        cr.font.color.rgb = _MUTED
                        cr.italic = True

                # --- Full active ads table (ALL of them, never truncated) ---
                all_ads = b.get("all_active_ads") or []
                if all_ads:
                    all_p = doc.add_paragraph()
                    all_p.add_run(f"Todos los anuncios activos ({b.get('all_active_ads_count', len(all_ads))})").bold = True

                    ads_table = doc.add_table(rows=1, cols=5)
                    ads_table.style = "Light Grid"
                    ads_table.autofit = False
                    col_widths = [Cm(2.0), Cm(2.2), Cm(2.4), Cm(3.0), Cm(7.0)]
                    for col, w in zip(ads_table.columns, col_widths):
                        col.width = w
                    hdr2 = ads_table.rows[0].cells
                    for i, label in enumerate(["Días activo", "Formato", "Hook", "Ángulo / Oferta", "Enlace (clic o copiar)"]):
                        hdr2[i].text = label
                        hdr2[i].width = col_widths[i]
                        _shade_cell(hdr2[i], _HEADER_SHADING)
                        for p in hdr2[i].paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.size = Pt(9)
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                    for i, ad in enumerate(all_ads):
                        row_cells = ads_table.add_row().cells
                        for c, w in zip(row_cells, col_widths):
                            c.width = w
                        if i % 2 == 1:
                            for c in row_cells:
                                _shade_cell(c, _ALT_ROW_SHADING)
                        row_cells[0].text = f"{ad['age_days']}d"
                        row_cells[1].text = _fmt_label(ad.get("format"))
                        row_cells[2].text = ad.get("hook_type") or "—"
                        angle_offer = " / ".join(
                            x for x in [ad.get("creative_angle"), (ad.get("offer_type") if ad.get("offer_type") != "no_offer" else None)] if x
                        ) or "—"
                        row_cells[3].text = angle_offer
                        link_para = row_cells[4].paragraphs[0]
                        ad_url = ad.get("ad_library_url")
                        # The link TEXT is the actual URL, not just "Ver ↗" — makes the
                        # link usable by copy-paste in any viewer that doesn't render OOXML
                        # hyperlinks as clickable, not only in ones that do.
                        _add_hyperlink(link_para, ad_url, ad_url or "no disponible")
                        for cell in row_cells[:4]:
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(8.5)
                        for p in row_cells[4].paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(7)

                doc.add_paragraph()  # spacer between brand cards

    # ---------- Footer / credit ----------
    doc.add_page_break()
    foot = doc.add_paragraph()
    foot.add_run(
        "Biblioteca generada automáticamente desde Meta Ad Library en vivo · Ecommerce Creative "
        "Intelligence · Señales de presencia publicitaria observable, no de ventas ni rentabilidad."
    ).italic = True
    foot2 = doc.add_paragraph()
    foot2.add_run(
        "Todos los anuncios enlazados son propiedad de sus respectivos anunciantes y se muestran a "
        "través de la Biblioteca de Anuncios de Meta (facebook.com/ads/library) — este documento no "
        "aloja ni redistribuye ese contenido, solo enlaza a la fuente oficial."
    ).font.color.rgb = _MUTED

    LIBRARY_DIR.mkdir(parents=True, exist_ok=True)
    path = LIBRARY_DIR / f"{niche}.docx"
    doc.save(str(path))
    return path
